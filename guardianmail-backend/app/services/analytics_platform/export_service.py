"""Multi-format export pipeline for analytics + reports.

Every generated report flows through here to be serialised into one of the
supported formats. Kept intentionally free of business logic — inputs are
plain dicts / dataclasses that came from `ReportingService`.
"""
from __future__ import annotations

import csv
import io
import json
from datetime import datetime
from typing import Any

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

from app.core.logging import get_logger

_log = get_logger(__name__)


class ExportService:
    # --------------------------------------------------------------- json
    def to_json(self, data: dict) -> tuple[bytes, str]:
        return (
            json.dumps(data, default=str, indent=2).encode(),
            "application/json",
        )

    # ---------------------------------------------------------------- csv
    def to_csv(self, rows: list[dict], *, header: list[str] | None = None) -> tuple[bytes, str]:
        if not rows and not header:
            return b"", "text/csv"
        header = header or list(rows[0].keys())
        buf = io.StringIO()
        w = csv.DictWriter(buf, fieldnames=header, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k) for k in header})
        return buf.getvalue().encode(), "text/csv"

    # --------------------------------------------------------------- xlsx
    def to_xlsx(self, sections: dict[str, list[dict]]) -> tuple[bytes, str]:
        wb = Workbook()
        first = True
        for sheet_name, rows in sections.items():
            ws = wb.active if first else wb.create_sheet()
            first = False
            ws.title = (sheet_name or "sheet")[:31]
            if not rows:
                ws.append(["(no data)"]); continue
            header = list(rows[0].keys())
            ws.append(header)
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", start_color="0F172A")
            for r in rows:
                ws.append([self._xl_value(r.get(h)) for h in header])
            for col_idx, h in enumerate(header, start=1):
                ws.column_dimensions[chr(64 + col_idx)].width = max(12, min(40, len(str(h)) + 4))
        out = io.BytesIO(); wb.save(out); out.seek(0)
        return out.getvalue(), (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    def _xl_value(self, v: Any) -> Any:
        if isinstance(v, (list, dict)):
            return json.dumps(v, default=str)
        return v

    # --------------------------------------------------------------- docx
    def to_docx(self, title: str, sections: list[dict]) -> tuple[bytes, str]:
        try:
            from docx import Document
        except Exception as exc:  # noqa: BLE001
            _log.error("docx_import_failed", error=str(exc))
            return self.to_json({"title": title, "sections": sections})
        doc = Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(f"Generated at {datetime.utcnow().isoformat()} UTC")
        for section in sections:
            doc.add_heading(section.get("title", "Section"), level=1)
            body = section.get("body")
            if isinstance(body, str):
                doc.add_paragraph(body)
            elif isinstance(body, list) and body and isinstance(body[0], dict):
                header = list(body[0].keys())
                table = doc.add_table(rows=1, cols=len(header))
                for i, h in enumerate(header):
                    table.rows[0].cells[i].text = str(h)
                for row in body:
                    cells = table.add_row().cells
                    for i, h in enumerate(header):
                        cells[i].text = str(row.get(h, ""))
        out = io.BytesIO(); doc.save(out); out.seek(0)
        return out.getvalue(), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # ---------------------------------------------------------------- pdf
    def to_pdf(self, title: str, sections: list[dict]) -> tuple[bytes, str]:
        try:
            from reportlab.lib.pagesizes import LETTER
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
            )
            from reportlab.lib import colors
        except Exception as exc:  # noqa: BLE001
            _log.error("reportlab_import_failed", error=str(exc))
            return self.to_json({"title": title, "sections": sections})

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=0.5 * inch,
                                rightMargin=0.5 * inch, topMargin=0.6 * inch,
                                bottomMargin=0.6 * inch)
        styles = getSampleStyleSheet()
        flow = [Paragraph(f"<b>{title}</b>", styles["Title"]),
                Paragraph(f"Generated at {datetime.utcnow().isoformat()} UTC",
                          styles["Italic"]),
                Spacer(1, 12)]
        for section in sections:
            flow.append(Paragraph(section.get("title", "Section"), styles["Heading2"]))
            body = section.get("body")
            if isinstance(body, str):
                flow.append(Paragraph(body, styles["BodyText"]))
            elif isinstance(body, list) and body and isinstance(body[0], dict):
                header = list(body[0].keys())
                data = [header] + [[str(r.get(h, "")) for h in header] for r in body[:100]]
                t = Table(data, hAlign="LEFT")
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ]))
                flow.append(t)
            flow.append(Spacer(1, 12))
        doc.build(flow)
        return buf.getvalue(), "application/pdf"
