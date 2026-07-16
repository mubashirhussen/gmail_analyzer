"""Evidence pack exporters (Module 9).

Renders an evidence bundle to the formats mandated by the module spec:

  * PDF  — human-readable summary (ReportLab)
  * DOCX — editable summary (python-docx)
  * JSON — canonical machine-readable manifest + integrity envelope
  * CSV  — indicator/URL table for spreadsheet triage
  * ZIP  — combined bundle containing all of the above + integrity file

The exporters take an already-assembled bundle dict (produced by the
platform_service) so they stay pure and unit-testable.
"""
from __future__ import annotations

import csv
import io
import json
import zipfile
from datetime import datetime
from typing import Any

from app.services.evidence.integrity import canonical_json


# --------------------------------------------------------------------------- #
# JSON                                                                        #
# --------------------------------------------------------------------------- #
def to_json(bundle: dict[str, Any]) -> bytes:
    return canonical_json(bundle)


# --------------------------------------------------------------------------- #
# CSV                                                                         #
# --------------------------------------------------------------------------- #
def to_csv(bundle: dict[str, Any]) -> bytes:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Section", "Field", "Value"])

    summary = bundle.get("summary", {})
    for k, v in summary.items():
        w.writerow(["summary", k, _stringify(v)])

    w.writerow([])
    w.writerow(["indicator_severity", "indicator_category", "indicator_detail"])
    for i in bundle.get("indicators", []):
        w.writerow([i.get("severity"), i.get("category"), i.get("detail")])

    w.writerow([])
    w.writerow(["url", "flagged", "providers"])
    for u in bundle.get("urls", []):
        w.writerow([
            u.get("url"),
            "yes" if u.get("flagged") else "no",
            ",".join(u.get("flagged_by", []) or []),
        ])

    w.writerow([])
    w.writerow(["attachment", "mime", "sha256"])
    for a in bundle.get("attachments", []):
        w.writerow([a.get("name"), a.get("mime"), a.get("sha256")])

    return buf.getvalue().encode("utf-8")


def _stringify(v: Any) -> str:
    if isinstance(v, (dict, list)):
        return json.dumps(v, default=str)
    return str(v) if v is not None else ""


# --------------------------------------------------------------------------- #
# DOCX                                                                        #
# --------------------------------------------------------------------------- #
def to_docx(bundle: dict[str, Any]) -> bytes:
    """Generate an editable DOCX summary using python-docx."""
    from docx import Document  # lazy import — heavy
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading("GuardianMail — Digital Evidence Pack", level=0)
    meta = bundle.get("integrity", {})
    doc.add_paragraph(f"Pack ID: {meta.get('pack_id')}")
    doc.add_paragraph(f"Generated: {meta.get('created_at')}")
    doc.add_paragraph(f"SHA-256: {meta.get('sha256')}")
    doc.add_paragraph(f"Signature (HMAC-SHA256): {meta.get('signature')}")

    _docx_section(doc, "Incident Summary", bundle.get("summary", {}))
    _docx_section(doc, "Message Metadata", bundle.get("message", {}))
    _docx_section(doc, "Authentication Results", bundle.get("authentication", {}))

    doc.add_heading("Threat Indicators", level=1)
    ind = bundle.get("indicators", [])
    if ind:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Severity"; hdr[1].text = "Category"; hdr[2].text = "Detail"
        for i in ind:
            row = table.add_row().cells
            row[0].text = str(i.get("severity", ""))
            row[1].text = str(i.get("category", ""))
            row[2].text = str(i.get("detail", ""))
    else:
        doc.add_paragraph("(no indicators recorded)")

    doc.add_heading("Suspicious URLs", level=1)
    urls = bundle.get("urls", [])
    if urls:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "URL"; hdr[1].text = "Flagged"; hdr[2].text = "Providers"
        for u in urls:
            row = table.add_row().cells
            row[0].text = str(u.get("url", ""))
            row[1].text = "yes" if u.get("flagged") else "no"
            row[2].text = ", ".join(u.get("flagged_by", []) or [])
    else:
        doc.add_paragraph("(no URLs recorded)")

    doc.add_heading("AI Analysis", level=1)
    ai = bundle.get("ai_report", {}) or {}
    doc.add_paragraph(ai.get("summary") or "(no AI summary available)")

    doc.add_heading("Timeline", level=1)
    for entry in bundle.get("timeline", []):
        doc.add_paragraph(f"{entry.get('at')} — {entry.get('event')}")

    doc.add_heading("Chain of Custody", level=1)
    for entry in bundle.get("chain_of_custody", []):
        doc.add_paragraph(
            f"{entry.get('at')} — {entry.get('event')} "
            f"(actor={entry.get('actor')})"
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_section(doc, title: str, mapping: dict[str, Any]) -> None:
    doc.add_heading(title, level=1)
    if not mapping:
        doc.add_paragraph("(none)")
        return
    for k, v in mapping.items():
        doc.add_paragraph(f"{k}: {_stringify(v)}")


# --------------------------------------------------------------------------- #
# PDF                                                                         #
# --------------------------------------------------------------------------- #
def to_pdf(bundle: dict[str, Any]) -> bytes:
    """Generate a printable PDF using ReportLab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, PageBreak,
    )
    from reportlab.lib import colors
    from reportlab.lib.units import inch

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
        title="GuardianMail Evidence Pack",
    )
    styles = getSampleStyleSheet()
    story: list = []

    story.append(Paragraph("GuardianMail — Digital Evidence Pack", styles["Title"]))
    meta = bundle.get("integrity", {})
    story.append(Paragraph(f"Pack ID: <b>{meta.get('pack_id')}</b>", styles["Normal"]))
    story.append(Paragraph(f"Generated: {meta.get('created_at')}", styles["Normal"]))
    story.append(Paragraph(f"SHA-256: <font face='Courier'>{meta.get('sha256')}</font>", styles["Normal"]))
    story.append(Paragraph(f"HMAC-SHA256 Signature: <font face='Courier'>{meta.get('signature')}</font>", styles["Normal"]))
    story.append(Spacer(1, 12))

    _pdf_kv(story, styles, "Incident Summary", bundle.get("summary", {}))
    _pdf_kv(story, styles, "Message", bundle.get("message", {}))
    _pdf_kv(story, styles, "Authentication", bundle.get("authentication", {}))

    story.append(Paragraph("Threat Indicators", styles["Heading2"]))
    ind_rows = [["Severity", "Category", "Detail"]]
    for i in bundle.get("indicators", []):
        ind_rows.append([str(i.get("severity", "")),
                         str(i.get("category", "")),
                         str(i.get("detail", ""))])
    story.append(_pdf_table(ind_rows))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Suspicious URLs", styles["Heading2"]))
    url_rows = [["URL", "Flagged", "Providers"]]
    for u in bundle.get("urls", []):
        url_rows.append([str(u.get("url", "")),
                         "yes" if u.get("flagged") else "no",
                         ", ".join(u.get("flagged_by", []) or [])])
    story.append(_pdf_table(url_rows))
    story.append(Spacer(1, 8))

    story.append(Paragraph("AI Analysis", styles["Heading2"]))
    ai_summary = (bundle.get("ai_report") or {}).get("summary") or "(none)"
    story.append(Paragraph(ai_summary, styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph("Chain of Custody", styles["Heading2"]))
    for entry in bundle.get("chain_of_custody", []):
        story.append(Paragraph(
            f"{entry.get('at')} — {entry.get('event')} "
            f"(actor={entry.get('actor')})",
            styles["Normal"],
        ))

    pdf.build(story)
    return buf.getvalue()


def _pdf_kv(story: list, styles, title: str, data: dict[str, Any]) -> None:
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    story.append(Paragraph(title, styles["Heading2"]))
    if not data:
        story.append(Paragraph("(none)", styles["Normal"]))
        story.append(Spacer(1, 6))
        return
    rows = [[k, _stringify(v)] for k, v in data.items()]
    tbl = Table(rows, hAlign="LEFT", colWidths=[130, None])
    tbl.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "Helvetica", 9),
        ("BOX", (0, 0), (-1, -1), 0.25, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BACKGROUND", (0, 0), (0, -1), colors.whitesmoke),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 8))


def _pdf_table(rows):
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    tbl = Table(rows, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "Helvetica", 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("BOX", (0, 0), (-1, -1), 0.25, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    return tbl


# --------------------------------------------------------------------------- #
# ZIP                                                                         #
# --------------------------------------------------------------------------- #
def to_zip(bundle: dict[str, Any]) -> bytes:
    """Bundle all supported formats + integrity metadata into a ZIP."""
    pack_id = (bundle.get("integrity") or {}).get("pack_id") or "pack"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{pack_id}/manifest.json", to_json(bundle))
        z.writestr(f"{pack_id}/indicators.csv", to_csv(bundle))
        try:
            z.writestr(f"{pack_id}/summary.pdf", to_pdf(bundle))
        except Exception as exc:  # pragma: no cover
            z.writestr(f"{pack_id}/summary.pdf.error.txt", str(exc))
        try:
            z.writestr(f"{pack_id}/summary.docx", to_docx(bundle))
        except Exception as exc:  # pragma: no cover
            z.writestr(f"{pack_id}/summary.docx.error.txt", str(exc))
        integrity = bundle.get("integrity", {})
        z.writestr(
            f"{pack_id}/INTEGRITY.txt",
            (
                f"pack_id  : {integrity.get('pack_id')}\n"
                f"version  : {integrity.get('version')}\n"
                f"created  : {integrity.get('created_at')}\n"
                f"sha256   : {integrity.get('sha256')}\n"
                f"hmac     : {integrity.get('signature')}\n"
                f"verify   : verify with the deployment SECRET_KEY.\n"
            ).encode("utf-8"),
        )
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# Dispatcher                                                                  #
# --------------------------------------------------------------------------- #
SUPPORTED_FORMATS = ("pdf", "docx", "json", "zip", "csv")

_MIME = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "json": "application/json",
    "zip": "application/zip",
    "csv": "text/csv",
}


def render(bundle: dict[str, Any], fmt: str) -> tuple[bytes, str]:
    fmt = fmt.lower()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"unsupported format {fmt}")
    match fmt:
        case "pdf":
            data = to_pdf(bundle)
        case "docx":
            data = to_docx(bundle)
        case "json":
            data = to_json(bundle)
        case "csv":
            data = to_csv(bundle)
        case "zip":
            data = to_zip(bundle)
        case _:  # pragma: no cover
            raise ValueError(fmt)
    return data, _MIME[fmt]
